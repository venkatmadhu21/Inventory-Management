import sqlite3
import docx       # this python library consists of all the modules,methods and classes needed to create and edit .docx files
from docx.shared import Inches  # Inches is a Class in shared module, which is included in the docx library. it used to mention dimensions in inches in word document.
from docx.enum.text import WD_ALIGN_PARAGRAPH #the text module consists of alsses and enumerations needed to modify paragraphs in word document, and we use WD_ALIGN_PARAGRAPH for alignment of paragraphs
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime


class Rv_form:
    def __init__(self):
        try:
            dbname = "Kmce1.db"
            con = sqlite3.connect(dbname)

            itemid="KMC1-IDC-1"
            supid="KMC1-MG-01"

            sql_query = f'''select i.itemname,s.sname from itemtable i,Suppliers s where s.supid='{supid}' and i.supid='{supid}' and i.itemid='{itemid}' '''
            data = con.execute(sql_query)
            result = data.fetchall()
            item_name=result[0][0]+"("+itemid+")"
            supplier_name=result[0][1]+"("+supid+")"

            amount=82000

            rv_no="RV-001-2024"
            doc_name=rv_no+".docx"

            document = docx.Document()  # a document object is created , this also creates an empty document in memory
            print("file opened successful...")
            heading_text = "Keshav Memorial Engineering College Of Engineering"
            heading = document.add_heading(heading_text,
                                           level=1)  # this statement calls add_heading function using document object to create an heading object giving text and heading kevel as parameters.
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # the heading is centre aligned using WD_ALIGN_PARAGRAPH.CENTER  property

            heading_text = "Koheda Road,Chinthapallyguda(V),Ibrahimpatnam(M),\nRR Dist-501510(T.S)Ph:9160102123,849981497"
            heading = document.add_heading(heading_text, level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            heading_text = "RECIEPT VOUCHER \n\n"
            heading = document.add_heading(heading_text, level=1)

            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            paragraph = document.add_paragraph()

            rv_date = datetime.date.today()
            fdate = rv_date.strftime("%d-%m-%y")
            paragraph.add_run(f'RV No: {rv_no}								 Date:{fdate}\n\n')

            table = document.add_table(rows=7,cols=4)  # the add_table method creates a table in the document with rows and columns as it's parameters

            fnames = table.rows[
                0].cells  # the row[] part allows us to access specific row of the table and the cells property gives access to each block in the table
            fnames[0].text = 'S.No'
            fnames[1].text = 'Item Name'
            fnames[2].text = 'Supplier Name'
            fnames[3].text = 'Amount'

            values = table.rows[1].cells
            values[0].text = '1.'
            values[1].text = f'{item_name}'
            values[2].text = f'{supplier_name}'
            values[3].text = f'{amount}'


            def set_cell_borders(
                    cell):  # this method creates borders around four corners of each cell making a visible table in word document

                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for side in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{side}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')  # Border size (1/8 pt)
                    border.set(qn('w:space'), '0')  # No space between borders
                    border.set(qn('w:color'), 'auto')  # Automatic color
                    tcBorders.append(border)
                tcPr.append(tcBorders)

            for row in table.rows:
                for cell in row.cells:
                    set_cell_borders(cell)

            paragraph1 = document.add_paragraph()

            paragraph1.add_run("\n\n\n\n\n\nRecieved By:")


            document.save(f'{doc_name}')

        except FileNotFoundError:
            print("error")
            return

if __name__ == "__main__":
    Rv_form()