import docx
from docx.shared import Inches,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class New_Form:
    def __init__(self):
        with open("IndentForm.txt","a+") as filo:
            ind_no=input("Enter Indent Number: ")
            date=input("Enter date(dd-mon-yyyy): ")
            dep=input("Enter name of the department: ")
            item_name=input("Enter Item name: ")
            supplier_name=input("enter supplier name: ")
            qty=int(input("Enter the quantity: "))
            price=int(input("enter price of the item: "))
            amount=price*qty
            empid=input("enter the employee id: ")
            approval=False
            approved_by="  "
            filo.write(f"{ind_no}\n{date}\n{dep}\n{item_name}\n{supplier_name}\n{qty}\n{price}\n{amount}\n{empid}\n{approval}\n{approved_by}")
            print("Indent form saved successfully....\n")


    def display_indent(i_no):
        with open("IndentForm.txt", "r") as filo:
            fields = filo.readlines()
            for count, i in enumerate(fields):
                if i.strip() == i_no:
                    break
            else:
                print("Indent number not found!!!")
                print()
                return 0

            fields1 = fields[count:count + 11]
            prompts = [ "Indent Number: ","Date (dd-mon-yyyy): ","Department Name: ","Item Name: ","Supplier Name: ","Quantity: ","Price per Item: ","Total Amount: ","Employee ID: ","Approval Status: ","Approved By: "]
            print(f"{'Indent form':^178}\n{'Keshav Memorial College of Engineering':^178}")
            for prompt, detail in zip(prompts, fields1):
                print(f"{prompt}{detail.strip()}")
            print()

class Edit_form(New_Form):
    def __init__(self):
        ind_no=input("enter the indent number: ")
        flag=New_Form.display_indent(ind_no)
        if flag!=0:
            with open("IndentForm.txt", "r+") as filo:
                fields = filo.readlines()
                while(1):
                    ov=input("Enter the value you want to change: ")
                    if ov=="false":
                        print("you are not authorised to approve the indent!!")
                        return
                    for i in fields:
                        if i.strip()==ov:
                            flag=1
                    if flag==1:
                        for count,k in enumerate(fields):
                            if k.strip()==ov:
                                nv = input("Enter the new value: ")+"\n"
                                fields[count]=nv
                                print(fields[count])
                                filo.seek(0)
                                filo.writelines(fields)
                                print("Value edited successfully...")
                                filo.close()
                                return
                    else:
                        print("value not found, Try again")
        else:
            return


class Create_file(New_Form):
    def __init__(self):
        ind_no=input("Enter the Indent Number: ")
        New_Form.display_indent(ind_no)
        ch=input("do you want to generate a Indent form(y/n): ")
        if ch=='y'or'Y':
            fnamed="IndentForm"+ind_no+".docx"
            with open("IndentForm.txt", "r") as filo1:
                fields = filo1.readlines()
                for count, i in enumerate(fields):
                    if i.strip() == ind_no:
                        flag=1
                        break
                if flag==1:
                    fields1 = fields[count:count + 11]
                    try:
                        document = docx.Document()
                        print("file opened successful...")
                        heading_text = "Keshav Memorial Engineering College Of Engineering"
                        heading = document.add_heading(heading_text, level=1)
                        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        heading_text = "Koheda Road,Chinthapallyguda(V),Ibrahimpatnam(M),\nRR Dist-501510(T.S)Ph:9160102123,849981497"
                        heading = document.add_heading(heading_text, level=3)
                        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        heading_text = "INDENT FORM\n\n"
                        heading = document.add_heading(heading_text, level=1)
                        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        paragraph = document.add_paragraph()


                        # Add 'Indent No:' and the indent number
                        paragraph.add_run(f'Indent No: {fields1[0].strip()}								         Date:{fields1[1]}\n\n')


                        paragraph.add_run(f'Name of the department: {fields1[2]}\n')

                        table = document.add_table(rows=7, cols=5)

                        fnames = table.rows[0].cells
                        fnames[0].text = 'S.No'
                        fnames[1].text = 'Item Name'
                        fnames[2].text = 'Supplier Name'
                        fnames[3].text = 'Quantity'
                        fnames[4].text = 'Total Cost'

                        values = table.rows[1].cells
                        values[0].text = '1.'
                        values[1].text = f'{fields1[3]}'
                        values[2].text = f'{fields1[4]}'
                        values[3].text = f'{fields1[5]}'
                        values[4].text = f'{fields1[7]}'


                        def set_cell_borders(cell):

                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            tcBorders = OxmlElement('w:tcBorders')
                            for side in ['top', 'left', 'bottom', 'right']:
                                border = OxmlElement(f'w:{side}')
                                border.set(qn('w:val'), 'single')
                                border.set(qn('w:sz'), '4')  # Border size (1/8 pt)
                                border.set(qn('w:space'), '0')  # No space between borders
                                border.set(qn('w:color'), 'auto')  # Automatic color
                                tcBorders.append(border)
                            tcPr.append(tcBorders)

                        for row in table.rows:
                            for cell in row.cells:
                                set_cell_borders(cell)

                        paragraph1 = document.add_paragraph()

                        paragraph1.add_run(f'\n\nRequested By: {fields1[8]}\n')

                        paragraph1.add_run('Admin/HOD signature: \n\n')

                        paragraph1.add_run('Approved By: \n\n\n')

                        paragraph1.add_run('Principal:                              							Director:')

                        paragraph1.add_run(' ')

                        document.save(f'{fnamed}')
                        print("Docx file created successfully , Check in this file location")


                    except FileNotFoundError:
                        print("Error")
                        return



while(1):
    print("1.New indent form\n2.Edit an indent form\n3.View indent\n4.Create a Indent.docx file\n5.Exit")
    choice=int(input("Enter your choice: "))

    if choice == 1:
        New_Form()
    elif choice==2:
        Edit_form()
    elif choice==3:
        ind_no=input("enter indent number: ")
        New_Form.display_indent(ind_no)
    elif choice==4:
        Create_file()
    else:
        exit()


