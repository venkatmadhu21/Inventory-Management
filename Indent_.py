import sqlite3
import docx  # this python library consists of all the modules,methods and classes needed to create and edit .docx files
from docx.shared import \
    Inches  # Inches is a Class in shared module, which is included in the docx library. it used to mention dimensions in inches in word document.
from docx.enum.text import \
    WD_ALIGN_PARAGRAPH  #the text module consists of alsses and enumerations needed to modify paragraphs in word document, and we use WD_ALIGN_PARAGRAPH for alignment of paragraphs
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime


class Indent_form:
    def __init__(self):
        try:

            itemid = "KMC1-REC-1"
            supid = "KMC1-LB-01"
            depipd = 'AD-01'
            empid = ''

            indent_no = "IN-001-2024"
            fname = indent_no + ".docx"

            in_date = datetime.date.today()
            fdate = in_date.strftime("%d-%m-%y")

            dbname = "Kmce1.db"
            con = sqlite3.connect(dbname)

            sql_query = f'''select i.itemid,i.itemname,s.sname from itemtable i,suppliers s where s.supid='{supid}' and i.supid='{supid}' and i.itemid='{itemid}' '''
            data = con.execute(sql_query)
            result = data.fetchall()

            sql_query = f''' select itemid, sum(quantity) as total_quantity from requirement where itemid = '{itemid}' '''
            data = con.execute(sql_query)
            quantity = data.fetchall()

            sql_query = f'''select depname from department where depid='{depipd}' '''
            data = con.execute(sql_query)
            depname = data.fetchall()

            item_name = result[0][0] + "(" + itemid + ")"
            supplier_name = result[0][1] + "(" + supid + ")"

            document = docx.Document()  # a document object is created , this also creates an empty document in memory
            print("file opened successful...")
            heading_text = "Keshav Memorial Engineering College Of Engineering"
            heading = document.add_heading(heading_text,
                                           level=1)  #this statement calls add_heading function using document object to create an heading object giving text and heading kevel as parameters.
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  #the heading is centre aligned using WD_ALIGN_PARAGRAPH.CENTER  property

            heading_text = "Koheda Road,Chinthapallyguda(V),Ibrahimpatnam(M),\nRR Dist-501510(T.S)Ph:9160102123,849981497"
            heading = document.add_heading(heading_text, level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            heading_text = "INDENT FORM\n\n"
            heading = document.add_heading(heading_text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            paragraph = document.add_paragraph()  # this statements adds an paragrapgh to the document

            paragraph.add_run(f'Indent No: {indent_no}								  Date:{fdate}\n\n')

            paragraph.add_run(
                f'Name of the department: {depname[0][0]}\n')  # the add.run method adds text to the paragrapgh and run formats the text with default font style sixe and other properties

            table = document.add_table(rows=7,cols=5)  # the add_table method creates a table in the document with rows and columns as it's parameters

            fnames = table.rows[
                0].cells  # the row[] part allows us to access specific row of the table and the cells property gives access to each block in the table
            fnames[0].text = 'S.No'
            fnames[1].text = 'Item Name'
            fnames[2].text = 'Supplier Name'
            fnames[3].text = 'Quantity'
            fnames[4].text = 'Total Cost'

            values = table.rows[1].cells
            values[0].text = '1.'
            values[1].text = f'{item_name}'
            values[2].text = f'{supplier_name}'
            values[3].text = f'{quantity[0][1]}'
            values[4].text = ''

            def set_cell_borders(
                    cell):  #this method creates borders around four corners of each cell making a visible table in word document

                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for side in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{side}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')  # Border size (1/8 pt)
                    border.set(qn('w:space'), '0')  # No space between borders
                    border.set(qn('w:color'), 'auto')  # Automatic color
                    tcBorders.append(border)
                tcPr.append(tcBorders)

            for row in table.rows:
                for cell in row.cells:
                    set_cell_borders(cell)

            paragraph1 = document.add_paragraph()

            paragraph1.add_run(f'\n\nRequested By: {depipd[0]}\n\n')

            paragraph1.add_run('Admin/HOD signature: \n\n')

            paragraph1.add_run('Approved By: \n\n\n')

            paragraph1.add_run('Principal:                              							Director:')

            paragraph1.add_run(' ')

            document.save(f'{fname}')  # the save method saves the document with the name given
            print("Docx file created successfully , Check in this file location")


        except FileNotFoundError:  # the FileNotFoundError class detects file not found error if file is not found
            print("Error")
            return


if __name__ == "__main__":
    Indent_form()
